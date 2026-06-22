import io
from docx import Document
from processors.docx import extract_texts, reinsert_texts


def _make_docx(texts: list[str]) -> bytes:
    doc = Document()
    for t in texts:
        doc.add_paragraph(t)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_table(rows: list[list[str]]) -> bytes:
    doc = Document()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for ri, row in enumerate(rows):
        for ci, text in enumerate(row):
            table.cell(ri, ci).text = text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_nested_table(outer_text: str, inner_text: str) -> bytes:
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    cell = outer.cell(0, 0)
    cell.paragraphs[0].add_run(outer_text)
    inner = cell.add_table(rows=1, cols=1)
    inner.cell(0, 0).paragraphs[0].add_run(inner_text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_texts():
    data = _make_docx(["안녕하세요", "테스트입니다"])
    segments = extract_texts(data)
    assert [s["text"] for s in segments] == ["안녕하세요", "테스트입니다"]


def test_reinsert_texts():
    data = _make_docx(["안녕하세요"])
    segments = extract_texts(data)
    translated = ["Hello"]
    result = reinsert_texts(data, segments, translated)
    doc = Document(io.BytesIO(result))
    assert doc.paragraphs[0].text == "Hello"


def test_empty_runs_skipped():
    data = _make_docx([""])
    segments = extract_texts(data)
    assert all(s["text"] != "" for s in segments)


def test_extract_texts_table():
    data = _make_docx_with_table([["분배", "문제"], ["기술", "반도체 역할을 설명하세요"]])
    segments = extract_texts(data)
    texts = [s["text"] for s in segments]
    assert "분배" in texts
    assert "문제" in texts
    assert "기술" in texts
    assert "반도체 역할을 설명하세요" in texts


def test_reinsert_texts_table():
    data = _make_docx_with_table([["분배", "문제"], ["기술", "반도체 역할을 설명하세요"]])
    segments = extract_texts(data)
    translated = ["Category", "Question", "Technology", "Explain the role of semiconductors"]
    result = reinsert_texts(data, segments, translated)
    doc = Document(io.BytesIO(result))
    all_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_texts.append(cell.text)
    assert "Category" in all_texts
    assert "Question" in all_texts
    assert "Technology" in all_texts
    assert "Explain the role of semiconductors" in all_texts


def test_extract_texts_nested_table():
    data = _make_docx_with_nested_table("외부 셀", "중첩 셀")
    segments = extract_texts(data)
    texts = [s["text"] for s in segments]
    assert "외부 셀" in texts
    assert "중첩 셀" in texts


def test_reinsert_texts_nested_table():
    data = _make_docx_with_nested_table("외부 셀", "중첩 셀")
    segments = extract_texts(data)
    translated = ["Outer Cell", "Nested Cell"]
    result = reinsert_texts(data, segments, translated)
    doc = Document(io.BytesIO(result))
    all_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_texts.append(cell.text)
    assert any("Outer Cell" in t for t in all_texts)
