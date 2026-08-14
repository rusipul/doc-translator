import io
from docx import Document


def _iter_table_paragraphs(table, base_key: tuple):
    # Keep strong references to _tc proxies so lxml doesn't GC them.
    seen_tcs: set = set()
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if cell._tc in seen_tcs:
                continue
            seen_tcs.add(cell._tc)
            cell_key = (*base_key, ri, ci)
            for pi, para in enumerate(cell.paragraphs):
                yield para, (*cell_key, pi)
            for nti, nested_table in enumerate(cell.tables):
                yield from _iter_table_paragraphs(nested_table, (*cell_key, nti))


def _iter_paragraphs(doc: Document):
    for i, para in enumerate(doc.paragraphs):
        yield para, ("para", i)
    for ti, table in enumerate(doc.tables):
        yield from _iter_table_paragraphs(table, ("table", ti))


def extract_texts(file_bytes: bytes) -> list[dict]:
    doc = Document(io.BytesIO(file_bytes))
    segments = []
    for para, key in _iter_paragraphs(doc):
        text = "".join(run.text for run in para.runs)
        if text.strip():
            segments.append({"text": text, "key": key})
    return segments


def reinsert_texts(file_bytes: bytes, segments: list[dict], translated: list[str]) -> bytes:
    doc = Document(io.BytesIO(file_bytes))
    key_to_translation = {
        tuple(s["key"]): t for s, t in zip(segments, translated)
    }
    for para, key in _iter_paragraphs(doc):
        t = key_to_translation.get(tuple(key))
        if t is not None:
            runs = para.runs
            if runs:
                runs[0].text = t
                for run in runs[1:]:
                    run.text = ""
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
